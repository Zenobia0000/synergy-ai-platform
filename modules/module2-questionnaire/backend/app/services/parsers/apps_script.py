"""
Apps Script docx 解析器（供問卷一致性驗證用）。

主要函式：parse_apps_script_docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

# 從 addXxxItem() 鏈式呼叫中抽取 setTitle('...') 的內容
# 覆蓋多行 chained calls
_TITLE_RE = re.compile(
    r"""\.setTitle\(\s*['"](?P<title>[^'"]+)['"]\s*\)""",
    re.MULTILINE,
)

# addXxx item types（用來確認是在題目上下文中）
_ITEM_TYPES_RE = re.compile(
    r"add(?:Text|MultipleChoice|Checkbox|Date|Paragraph|Scale|Grid|Time|SectionHeader|PageBreak)Item\s*\(",
)


def _extract_full_text(docx_path: Path) -> str:
    """讀取 docx 所有段落與表格單元格，合併為完整字串。"""
    doc = Document(str(docx_path))
    lines: list[str] = []

    for para in doc.paragraphs:
        lines.append(para.text)

    # 若有 table，也納入
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    lines.append(para.text)

    return "\n".join(lines)


def parse_apps_script_docx(path: Path) -> tuple[str, ...]:
    """
    解析 Apps Script docx，抽取所有 setTitle() 呼叫的題目標題。

    策略：
    1. 讀取所有段落合併為完整文字
    2. 正則抽取所有 .setTitle('...')
    3. 過濾空白，去重並保持出現順序
    4. 返回 tuple[str, ...]

    Args:
        path: Apps Script docx 路徑

    Returns:
        題目標題 tuple（依文件順序，去重後）

    Raises:
        FileNotFoundError: 檔案不存在
        ValueError: 無法解析文件
    """
    if not path.exists():
        raise FileNotFoundError(f"Apps Script docx 不存在：{path}")

    full_text = _extract_full_text(path)

    if not full_text.strip():
        raise ValueError(f"docx 內容為空：{path}")

    # 抽取所有 setTitle 內容（保序去重）
    seen: set[str] = set()
    titles: list[str] = []

    for m in _TITLE_RE.finditer(full_text):
        title = m.group("title").strip()
        if not title:
            continue
        if title not in seen:
            seen.add(title)
            titles.append(title)

    return tuple(titles)
